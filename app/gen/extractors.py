from __future__ import annotations

from typing import Tuple

import re


def _clean_text(raw: str) -> str:
    if not raw:
        return ""
    t = raw
    # Normalize HTML-ish breaks and strip tags
    t = t.replace("<br>", "\n").replace("<br/>", "\n").replace("<br />", "\n")
    t = re.sub(r"<[^>]+>", " ", t)
    # Normalize hyphenated line-breaks and whitespace
    t = t.replace("\r", "\n")
    t = re.sub(r"(\w)-\n(\w)", r"\1\2", t)
    t = re.sub(r"\n{3,}", "\n\n", t)
    t = re.sub(r"[\t\x0b\x0c]+", " ", t)
    t = re.sub(r" +", " ", t)
    return t.strip()


def extract_text_from_txt_bytes(data: bytes) -> str:
    try:
        txt = data.decode("utf-8", errors="replace")
    except Exception:
        try:
            txt = data.decode("latin-1", errors="replace")
        except Exception:
            txt = ""
    return _clean_text(txt)


def extract_text_from_docx_bytes(data: bytes) -> str:
    try:
        import io
        import docx  # type: ignore

        d = docx.Document(io.BytesIO(data))
        parts: list[str] = []
        for p in d.paragraphs:
            if p.text and p.text.strip():
                parts.append(p.text.strip())
        for tbl in d.tables:
            for row in tbl.rows:
                parts.append(" | ".join((c.text or "").strip() for c in row.cells))
        return _clean_text("\n".join(parts))
    except Exception:
        return ""


def extract_text_from_pdf_bytes(data: bytes, *, poppler_path: str | None = None, ocr_dpi: int = 300, ocr_max_pages: int = 30, max_chars: int = 8000) -> str:
    from io import BytesIO

    chunks: list[str] = []
    # 1) pdfplumber: page text + tables
    try:
        import pdfplumber  # type: ignore

        with pdfplumber.open(BytesIO(data)) as pdf:
            for page in pdf.pages:
                t = (page.extract_text() or "").strip()
                if t:
                    chunks.append(t)
                try:
                    tables = page.extract_tables() or []
                    for tbl in tables:
                        rows = [" | ".join((cell or "").strip() for cell in row) for row in (tbl or [])]
                        if rows:
                            chunks.append("\n".join(rows))
                except Exception:
                    pass
                if len(_clean_text("\n\n".join(chunks))) >= max_chars - 200:
                    break
    except Exception:
        pass

    text_content = "\n\n".join(chunks)

    # 2) pdfminer fallback (some PDFs extract better this way)
    try:
        import pdfminer.high_level as pdf_hi  # type: ignore

        mined = pdf_hi.extract_text(BytesIO(data)) or ""
        if mined:
            text_content = (text_content + "\n\n" + mined) if text_content else mined
    except Exception:
        pass

    cleaned = _clean_text(text_content)
    if len(cleaned) >= max_chars or len(cleaned) >= 200:
        return cleaned[:max_chars]

    # 3) OCR (images / scanned PDFs)
    try:
        from pdf2image import convert_from_bytes  # type: ignore
        import pytesseract  # type: ignore

        images = convert_from_bytes(data, poppler_path=poppler_path, dpi=ocr_dpi)
        ocr_chunks: list[str] = []
        for img in images[:ocr_max_pages]:
            try:
                ocr_chunks.append(pytesseract.image_to_string(img) or "")
            except Exception:
                continue
            combined = _clean_text(((text_content + "\n\n") if text_content else "") + "\n\n".join(ocr_chunks))
            if len(combined) >= max_chars - 50:
                cleaned = combined
                break
        else:
            cleaned = _clean_text(((text_content + "\n\n") if text_content else "") + "\n\n".join(ocr_chunks))
        return cleaned[:max_chars]
    except Exception:
        return cleaned[:max_chars]


def extract_text_from_upload(file_storage, *, poppler_path: str | None = None, ocr_dpi: int = 300, ocr_max_pages: int = 30, max_chars: int = 8000) -> Tuple[str, str]:
    """Return (clean_text, ext). ext is lowercase extension without dot."""
    filename = getattr(file_storage, "filename", "") or ""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    data = file_storage.read() or b""
    if ext == "txt":
        return extract_text_from_txt_bytes(data)[:max_chars], ext
    if ext == "docx":
        return extract_text_from_docx_bytes(data)[:max_chars], ext
    if ext == "pdf":
        return extract_text_from_pdf_bytes(data, poppler_path=poppler_path, ocr_dpi=ocr_dpi, ocr_max_pages=ocr_max_pages, max_chars=max_chars), ext
    # Fallback try text decode
    return extract_text_from_txt_bytes(data)[:max_chars], ext


